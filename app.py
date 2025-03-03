import streamlit as st
import os
from PyPDF2 import PdfReader
from docx import Document
import google.generativeai as genai
import re
from io import BytesIO

# Page configuration
st.set_page_config(page_title="Advanced Document Analysis Assistant", layout="wide")

def initialize_ai_api():
    """Initialize the AI API safely"""
    try:
        # Check if API key is in Streamlit secrets
        if "GEMINI_API_KEY" in st.secrets:
            api_key = st.secrets["GEMINI_API_KEY"]
        else:
            st.warning("⚠️ API Key not found in secrets!")
            
            # Ask for API key input
            api_key = st.text_input(
                "Enter your AI API Key (free to create at https://aistudio.google.com/)",
                type="password"
            )
            
            if not api_key:
                st.info("You need a Google AI Studio account to get a free API key.")
                st.info("1. Visit https://aistudio.google.com/")
                st.info("2. Create a free account")
                st.info("3. Get your API key from the settings")
                return None
        
        # Configure the API
        genai.configure(api_key=api_key)
        
        # Set the model
        model_name = "models/gemini-2.0-flash"
        
        # Save key to session state for current session only
        st.session_state.temp_api_key = api_key
        st.session_state.selected_model = model_name
        
        st.success(f"Advanced AI model initialized successfully! Ready to analyze your documents.")
        
        # Create the model instance
        return genai.GenerativeModel(model_name)
            
    except Exception as e:
        st.error(f"Error initializing AI API: {str(e)}")
        return None

def extract_text_from_pdf(file):
    """Extract text from PDF with improved handling"""
    try:
        reader = PdfReader(file)
        
        # Extract text page by page with formatting preservation
        text_content = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                # Add page markers for better context
                text_content.append(f"--- Page {i+1} ---\n{page_text}")
        
        full_text = "\n\n".join(text_content)
        
        # If no text was extracted, it might be a scanned PDF
        if not full_text.strip():
            st.warning("The PDF appears to be scanned or doesn't contain extractable text.")
            return None
            
        return full_text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
        return None

def extract_text_from_docx(file):
    """Extract text from DOCX with improved structure preservation"""
    try:
        doc = Document(BytesIO(file.getvalue()))
        
        # Extract structured content
        content = []
        
        # Extract headers and paragraphs with formatting
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                # Preserve heading structure
                content.append(f"## {para.text}")
            elif para.text.strip():
                content.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
            content.append("\nTable:\n" + "\n".join(table_text) + "\n")
            
        return "\n".join(content)
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {str(e)}")
        return None

def extract_text(file):
    """Extract text from PDF, DOCX, or TXT files with improved handling"""
    try:
        if file.type == "application/pdf":
            return extract_text_from_pdf(file)
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return extract_text_from_docx(file)
        elif file.type == "text/plain":
            text = file.read().decode("utf-8", errors="replace")
            return text
        else:
            st.error(f"Unsupported file type: {file.type}")
            return None
    except Exception as e:
        st.error(f"Error extracting text: {str(e)}")
        return None

def preprocess_document(text):
    """Enhance document structure for better AI understanding"""
    if not text:
        return text
        
    # Clean up extra whitespace while preserving paragraph breaks
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Identify and mark code blocks (for technical docs)
    code_pattern = r'(?:proc\s+\w+|data\s+\w+|libname\s+\w+|\w+\s*=\s*\w+\(|\w+\s*<-\s*\w+\()'
    lines = text.split('\n')
    in_code_block = False
    
    for i, line in enumerate(lines):
        if re.search(code_pattern, line, re.IGNORECASE) and not in_code_block:
            lines[i] = "```\n" + line
            in_code_block = True
        elif in_code_block and line.strip() == '':
            if i > 0:
                lines[i-1] = lines[i-1] + "\n```"
            in_code_block = False
    
    # Close any open code block
    if in_code_block:
        lines.append("```")
    
    return '\n'.join(lines)

def identify_document_type(text):
    """Identify document type based on content to improve context"""
    if not text:
        return "Unknown"
        
    # Check for research paper indicators
    research_indicators = ["abstract", "introduction", "methodology", "results", "discussion", "conclusion", "references"]
    research_score = sum(1 for indicator in research_indicators if re.search(r'\b' + indicator + r'\b', text.lower()))
    
    # Check for technical document indicators
    technical_indicators = ["proc", "function", "method", "class", "library", "import", "code", "algorithm"]
    technical_score = sum(1 for indicator in technical_indicators if re.search(r'\b' + indicator + r'\b', text.lower()))
    
    # Check for statistical content
    statistical_indicators = ["mean", "median", "standard deviation", "variance", "correlation", "regression", "p-value", "hypothesis"]
    statistical_score = sum(1 for indicator in statistical_indicators if re.search(r'\b' + indicator + r'\b', text.lower()))
    
    # Determine document type
    if research_score >= 3:
        return "Research Paper"
    elif technical_score >= 3:
        if statistical_score >= 2:
            return "Statistical/Technical Document"
        return "Technical Document"
    elif statistical_score >= 3:
        return "Statistical Document"
    else:
        return "General Document"

def find_relevant_chunk(context, question, chunk_size=12000, overlap=2000):
    """Find the most relevant chunk with keyword-based scoring and section awareness"""
    if not context or len(context) <= chunk_size:
        return context
    
    # Extract keywords from question
    question_words = set(re.findall(r'\b\w+\b', question.lower()))
    question_words = {w for w in question_words if len(w) > 3}  # Filter out short words
    
    # Create chunks with overlap
    chunks = []
    start = 0
    while start < len(context):
        end = min(start + chunk_size, len(context))
        
        # Try to find section boundaries for cleaner chunks
        if end < len(context):
            # Look for section breaks within 500 chars of the end point
            section_boundary = context.find("\n##", end - 500, end + 500)
            if section_boundary != -1:
                end = section_boundary
                
        chunks.append(context[start:end])
        start += chunk_size - overlap
    
    # Score chunks based on keyword matching and position
    chunk_scores = []
    
    for i, chunk in enumerate(chunks):
        # Count keyword matches
        keyword_score = sum(10 for word in question_words if word in chunk.lower())
        
        # Check for exact phrases (higher weight)
        for phrase in re.findall(r'\b\w+(?:\s+\w+){1,5}\b', question.lower()):
            if phrase in chunk.lower():
                keyword_score += 15
        
        # Position bias - favor earlier chunks
        position_score = max(5 - i, 0) if i < 5 else 0
        
        # Check for section headings that match keywords
        section_bonus = 0
        for heading in re.findall(r'##\s*(.*?)(?=\n|$)', chunk):
            heading_lower = heading.lower()
            if any(word in heading_lower for word in question_words):
                section_bonus += 20
        
        # Code block bonus
        code_block_bonus = 0
        if "```" in chunk and any(tech_word in question.lower() for tech_word in ["code", "function", "proc", "example"]):
            code_block_bonus = 15
            
        total_score = keyword_score + position_score + section_bonus + code_block_bonus
        chunk_scores.append(total_score)
    
    # If no good matches, return the first chunk as fallback
    if max(chunk_scores) <= 0:
        return chunks[0]
        
    # Get best chunk
    best_chunk_index = chunk_scores.index(max(chunk_scores))
    
    # If the score is very high, just return this chunk
    if chunk_scores[best_chunk_index] > 50:
        return chunks[best_chunk_index]
    
    # For moderate scores, combine with adjacent chunks if they also have good scores
    if chunk_scores[best_chunk_index] > 10:
        combined_text = chunks[best_chunk_index]
        
        # Check if previous chunk is also relevant
        if best_chunk_index > 0 and chunk_scores[best_chunk_index-1] > 5:
            combined_text = chunks[best_chunk_index-1] + "\n\n" + combined_text
            
        # Check if next chunk is also relevant
        if best_chunk_index < len(chunks)-1 and chunk_scores[best_chunk_index+1] > 5:
            combined_text = combined_text + "\n\n" + chunks[best_chunk_index+1]
            
        # If combined text is too long, trim it
        if len(combined_text) > chunk_size * 2:
            combined_text = combined_text[:chunk_size * 2]
            
        return combined_text
    
    return chunks[best_chunk_index]

def generate_response(model, question, context, doc_type):
    """Generate answer using AI with improved prompting based on document type"""
    if not model:
        return "API configuration error. Please check your API key and model selection."
    
    # Find relevant chunk to fit within model context limits
    relevant_context = find_relevant_chunk(context, question)
    
    # Create specific prompts based on document type
    base_prompt = f"""You are an expert document analysis assistant. A user has uploaded a {doc_type}. 
Answer the following question based ONLY on the provided document content.
If the information is not in the document, say 'I don't see that information in the document.' 
Do not make up information or rely on external knowledge.

Document Content:
{relevant_context}

Question: {question}
"""
    
    # Add specific instructions based on document type
    if "Research Paper" in doc_type:
        prompt = base_prompt + """
Provide a thorough academic answer. Include:
1. Direct quotes or specific sections from the paper when relevant
2. Methodological details if asked about methods
3. Clear explanations of findings and their implications
4. Precise terminology used in the paper

If asked for a summary, provide a structured summary with Background, Methods, Results, and Conclusions.
"""
    elif "Statistical" in doc_type:
        prompt = base_prompt + """
Provide a detailed technical answer with:
1. Precise statistical terminology and definitions from the document
2. Complete code examples if present in the document and relevant to the question
3. Step-by-step explanations of statistical procedures
4. Interpretations of statistical output if present

If asked about a statistical procedure or programming code, provide:
- Full syntax/code examples if found in the document
- Detailed explanation of parameters and options
- Context about when and how to use this procedure
- Example output interpretation if available
"""
    elif "Technical Document" in doc_type:
        prompt = base_prompt + """
Provide a technical answer with:
1. Code examples and syntax from the document when relevant
2. Precise technical terminology as used in the document
3. Step-by-step instructions if asked for a procedure
4. Links between concepts if they appear in different sections

For code questions, show complete examples from the document rather than fragments.
"""
    else:
        prompt = base_prompt + """
Provide a clear, organized response that directly addresses the question.
Include specific information from the document, and organize complex answers with headings or bullet points if needed.
"""

    try:
        generation_config = {
            "temperature": 0.2,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 1024,
        }
        
        response = model.generate_content(prompt, generation_config=generation_config)
        return response.text
    except Exception as e:
        return f"Error generating response: {str(e)}"

def main():
    st.title("📄 Advanced Document Analysis Assistant")
    st.write("Upload a document and ask detailed questions about its content!")
    
    # Initialize the model
    model = None
    if "temp_api_key" in st.session_state and "selected_model" in st.session_state:
        genai.configure(api_key=st.session_state.temp_api_key)
        model = genai.GenerativeModel(st.session_state.selected_model)
        st.success(f"Advanced AI model initialized successfully! Ready to analyze your documents.")
    else:
        model = initialize_ai_api()
    
    # File uploader
    uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "txt"])
    
    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    # Initialize document context
    if "context" not in st.session_state:
        st.session_state.context = None
        st.session_state.doc_type = None
    
    # Process uploaded file
    if uploaded_file:
        col1, col2 = st.columns([3, 1])
        with col1:
            if st.button("📊 Process Document") or st.session_state.context is None:
                with st.spinner("Processing document..."):
                    # Extract text
                    raw_context = extract_text(uploaded_file)
                    
                    if raw_context:
                        # Preprocess the document to enhance structure
                        processed_context = preprocess_document(raw_context)
                        
                        # Identify document type
                        doc_type = identify_document_type(processed_context)
                        
                        # Save to session state
                        st.session_state.context = processed_context
                        st.session_state.doc_type = doc_type
                        st.session_state.file_name = uploaded_file.name
                        
                        # Display document info
                        doc_size_kb = len(processed_context) / 1024
                        token_estimate = len(processed_context) / 4  # Rough estimate
                        
                        st.success(f"✅ Document processed successfully!")
                        st.info(f"📄 Document type: {doc_type}")
                        st.info(f"📏 Size: ~{doc_size_kb:.1f} KB, ~{token_estimate:.0f} tokens")
        
        with col2:
            if st.session_state.context is not None:
                if st.button("🔄 Clear Chat History"):
                    st.session_state.messages = []
                    st.experimental_rerun()
    
    # Display chat interface if document is processed
    if st.session_state.get("context"):
        st.write(f"💬 Analyzing: **{st.session_state.get('file_name', 'Document')}** ({st.session_state.get('doc_type', 'Document')})")
        
        # Display chat history
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # Handle user input
        if question := st.chat_input("Ask any question about your document..."):
            # Add user message to chat history
            st.session_state.messages.append({"role": "user", "content": question})
            
            # Display user message
            with st.chat_message("user"):
                st.markdown(question)
            
            # Generate and display assistant response
            with st.spinner("Analyzing document for the best answer..."):
                answer = generate_response(model, question, st.session_state.context, st.session_state.doc_type)
            
            # Add assistant response to chat history
            st.session_state.messages.append({"role": "assistant", "content": answer})
            
            # Display assistant response
            with st.chat_message("assistant"):
                st.markdown(answer)

if __name__ == "__main__":
    main()
